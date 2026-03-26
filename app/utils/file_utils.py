"""파일 처리 관련 유틸리티 — 텍스트 추출 + 업로드 검증"""

import logging
import os
from pathlib import Path
from typing import Union

from PyPDF2 import PdfReader

from app.config import settings
from app.exceptions import FileFormatError, FileProcessingError, FileSizeExceededError, InvalidRequestError

logger = logging.getLogger(__name__)


def validate_file_type(filename: str) -> bool:
    """
    파일 확장자 검증

    Args:
        filename: 파일명

    Returns:
        허용된 파일 타입이면 True
    """
    suffix = Path(filename).suffix.lower()
    return suffix in settings.allowed_file_extensions


def extract_text_from_pdf(file_path: Union[str, Path]) -> str:
    """
    PDF 파일에서 텍스트 추출

    Args:
        file_path: PDF 파일 경로

    Returns:
        추출된 텍스트

    Raises:
        FileProcessingError: PDF 읽기 실패 시
    """
    try:
        reader = PdfReader(str(file_path))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF 텍스트 추출 실패: {e}")
        raise FileProcessingError(
            f"PDF 파일을 읽을 수 없습니다: {file_path}",
            details={"path": str(file_path), "error": str(e)}
        )


def extract_text_from_docx(file_path: Union[str, Path]) -> str:
    """
    DOCX 파일에서 텍스트 추출

    Args:
        file_path: DOCX 파일 경로

    Returns:
        추출된 텍스트

    Raises:
        FileProcessingError: DOCX 읽기 실패 시
    """
    try:
        from docx import Document

        doc = Document(str(file_path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        logger.error(f"DOCX 텍스트 추출 실패: {e}")
        raise FileProcessingError(
            f"DOCX 파일을 읽을 수 없습니다: {file_path}",
            details={"path": str(file_path), "error": str(e)}
        )


def extract_text_from_hwp(file_path: Union[str, Path]) -> str:
    """
    HWP 파일에서 텍스트 추출

    시그니처 기반 자동 감지:
    - OLE2 (진짜 HWP v5): olefile로 PrvText 스트림 추출
    - ZIP (확장자만 .hwp인 HWPX): python-hwpx로 추출

    Args:
        file_path: HWP 파일 경로

    Returns:
        추출된 텍스트
    """
    file_path = Path(file_path)
    try:
        header = file_path.read_bytes()[:4]
    except Exception as e:
        raise FileProcessingError(
            f"HWP 파일을 읽을 수 없습니다: {file_path}",
            details={"path": str(file_path), "error": str(e)}
        )

    # ZIP 시그니처 (PK) → 실제로는 HWPX
    if header[:2] == b"PK":
        logger.info(f"HWP 파일이 실제로는 HWPX (ZIP): {file_path.name}")
        try:
            from hwpx import HwpxDocument
            doc = HwpxDocument(str(file_path))
            text = doc.to_text() if hasattr(doc, "to_text") else ""
            if not text:
                from hwpx import TextExtractor
                extractor = TextExtractor(str(file_path))
                text = extractor.extract()
            return text
        except Exception as e:
            logger.warning(f"HWPX 파싱 실패, 빈 문자열 반환: {e}")
            return ""

    # OLE2 시그니처 (D0 CF 11 E0) → 진짜 HWP v5
    if header[:4] == b"\xd0\xcf\x11\xe0":
        try:
            import olefile
            ole = olefile.OleFileIO(str(file_path))
            try:
                if ole.exists("PrvText"):
                    raw = ole.openstream("PrvText").read()
                    return raw.decode("utf-16")
                else:
                    logger.warning(f"HWP PrvText 스트림 없음: {file_path.name}")
                    return ""
            finally:
                ole.close()
        except Exception as e:
            logger.warning(f"HWP OLE2 파싱 실패: {e}")
            return ""

    logger.warning(f"HWP 파일 시그니처 불명: {header.hex()}")
    return ""


def extract_text_from_file(file_path: Union[str, Path]) -> str:
    """
    파일에서 텍스트 추출 (통합 함수)

    Args:
        file_path: 파일 경로

    Returns:
        추출된 텍스트

    Raises:
        FileProcessingError: 파일 읽기 실패 또는 지원하지 않는 형식
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileProcessingError(
            f"파일을 찾을 수 없습니다: {file_path}",
            details={"path": str(file_path)}
        )

    suffix = file_path.suffix.lower()

    try:
        if suffix == ".pdf":
            return extract_text_from_pdf(file_path)
        elif suffix == ".docx":
            return extract_text_from_docx(file_path)
        elif suffix == ".hwp":
            return extract_text_from_hwp(file_path)
        elif suffix == ".txt":
            return file_path.read_text(encoding="utf-8")
        else:
            raise FileProcessingError(
                f"지원하지 않는 파일 형식입니다: {suffix}",
                details={"path": str(file_path), "suffix": suffix}
            )
    except FileProcessingError:
        raise
    except Exception as e:
        logger.error(f"파일 텍스트 추출 실패: {e}")
        raise FileProcessingError(
            f"파일을 읽는 중 오류가 발생했습니다: {file_path}",
            details={"path": str(file_path), "error": str(e)}
        )


# ── 업로드 검증 ──────────────────────────────────────────────

# 카테고리별 허용 확장자 (점 없이 소문자)
UPLOAD_ALLOWED_EXTENSIONS: dict[str, frozenset[str]] = {
    "project_file": frozenset({
        "pdf", "docx", "hwp", "hwpx", "xlsx", "pptx", "png", "jpg", "jpeg",
    }),
    "asset": frozenset({"pdf", "docx"}),
    "template": frozenset({"pdf", "docx"}),
    "submission": frozenset({
        "pdf", "hwp", "hwpx", "doc", "docx", "xls", "xlsx",
        "ppt", "pptx", "jpg", "jpeg", "png", "zip", "txt",
    }),
}


def sanitize_filename(filename: str | None) -> str:
    """파일명 보안 검증 (경로 탐색 방지). 안전한 파일명 반환."""
    raw = filename or "upload"
    safe = os.path.basename(raw).strip()
    if not safe or safe.startswith("."):
        raise InvalidRequestError("잘못된 파일명입니다.")
    return safe


def get_extension(filename: str) -> str:
    """파일명에서 확장자 추출 (소문자, 점 없이)."""
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def validate_extension(filename: str, category: str) -> str:
    """확장자 검증. 유효하면 확장자 반환, 아니면 FileFormatError."""
    allowed = UPLOAD_ALLOWED_EXTENSIONS.get(category)
    if allowed is None:
        raise ValueError(f"알 수 없는 파일 카테고리: {category}")
    ext = get_extension(filename)
    if ext not in allowed:
        raise FileFormatError(ext, sorted(allowed))
    return ext


def validate_file_size(content: bytes, max_mb: int | None = None):
    """파일 크기 검증. 초과 시 FileSizeExceededError."""
    limit = max_mb or settings.max_file_size_mb
    max_bytes = limit * 1024 * 1024
    if len(content) > max_bytes:
        raise FileSizeExceededError(limit, len(content) / (1024 * 1024))


def validate_upload(filename: str | None, content: bytes, category: str, max_mb: int | None = None) -> tuple[str, str]:
    """파일명 + 확장자 + 크기 통합 검증. (safe_filename, extension) 반환."""
    safe = sanitize_filename(filename)
    ext = validate_extension(safe, category)
    validate_file_size(content, max_mb)
    return safe, ext
