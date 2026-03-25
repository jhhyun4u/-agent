"""
산출내역서 DOCX 빌더

bid_plan(STEP 2.5) + plan_price(STEP 3)의 원가 데이터를 기반으로
공공조달 표준 양식의 산출내역서 문서를 생성한다.

■ 문서 구조
1. 표지 (사업명, 제출자, 날짜)
2. 총괄표 (비목별 합계)
3. 직접인건비 상세 (등급별 MM × 단가)
4. 직접경비 상세 (항목별)
5. 간접비·기술료·이윤
6. 산출 근거 (Budget Narrative)
"""

import logging
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)


def _fmt_won(amount: int | float) -> str:
    """금액 포맷팅 (천원 단위 쉼표)."""
    return f"{int(amount):,}"


def _set_cell(cell, text: str, bold: bool = False, align: str = "left", size: int = 9):
    """셀 텍스트 + 스타일 설정."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = "맑은 고딕"
    run.bold = bold
    if align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _shade_row(row, color: str = "F2F2F2"):
    """행 배경색 설정."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in row.cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)


def build_cost_sheet(
    project_name: str,
    client: str,
    proposer_name: str,
    bid_plan_data: dict | None,
    plan_price_data: dict | None,
    cost_standard: str = "KOSA",
    year: int | None = None,
) -> BytesIO:
    """산출내역서 DOCX 생성.

    Args:
        project_name: 사업명
        client: 발주기관
        proposer_name: 제안업체명
        bid_plan_data: bid_plan state (BidPlanResult.model_dump())
        plan_price_data: plan.bid_price dict (plan_price 노드 산출물)
        cost_standard: 원가 기준 (KOSA/KEA/MOEF)
        year: 적용 연도

    Returns:
        BytesIO: DOCX 파일 바이트 스트림
    """
    doc = Document()
    year = year or datetime.now().year

    # 기본 스타일
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    bid_plan = bid_plan_data or {}
    plan_price = plan_price_data or {}

    # ── 데이터 추출 ──
    cost_breakdown = bid_plan.get("cost_breakdown", {})
    labor_cost = plan_price.get("labor_cost", {})
    labor_breakdown = labor_cost.get("breakdown", [])
    labor_total = labor_cost.get("total", cost_breakdown.get("direct_labor", 0))

    direct_expenses = plan_price.get("direct_expenses", {})
    expense_items = direct_expenses.get("items", [])
    expense_total = direct_expenses.get("total", 0)

    overhead = plan_price.get("overhead", {})
    overhead_rate = overhead.get("rate", 1.10) if isinstance(overhead, dict) else 1.10
    overhead_total = overhead.get("total", 0) if isinstance(overhead, dict) else 0

    profit = plan_price.get("profit", {})
    profit_rate = profit.get("rate", 0.22) if isinstance(profit, dict) else 0.22
    profit_total = profit.get("total", 0) if isinstance(profit, dict) else 0

    total_cost = plan_price.get("total_cost", bid_plan.get("recommended_bid", 0))

    budget_narrative = plan_price.get("budget_narrative", [])

    # 간접비가 0이면 직접인건비 기반 추정
    if not overhead_total and labor_total:
        overhead_total = int(labor_total * (overhead_rate if overhead_rate > 1 else overhead_rate * labor_total))
        if overhead_rate > 1:
            overhead_total = int(labor_total * (overhead_rate - 1))

    # 기술료가 0이면 추정
    if not profit_total and labor_total:
        profit_total = int((labor_total + overhead_total) * profit_rate)

    subtotal = labor_total + expense_total + overhead_total + profit_total
    vat = int(subtotal * 0.10)
    grand_total = total_cost if total_cost else subtotal + vat

    # ═══════════════════════════════════════
    # 1. 표지
    # ═══════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("산 출 내 역 서")
    run.font.size = Pt(28)
    run.bold = True

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(project_name)
    run.font.size = Pt(16)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"\n발주기관: {client}").font.size = Pt(12)
    meta.add_run(f"\n\n제 출 자: {proposer_name}").font.size = Pt(12)
    meta.add_run(f"\n\n{datetime.now().strftime('%Y년 %m월 %d일')}").font.size = Pt(12)

    doc.add_page_break()

    # ═══════════════════════════════════════
    # 2. 총괄표
    # ═══════════════════════════════════════
    h = doc.add_heading("1. 사업비 총괄표", level=1)
    h.runs[0].font.size = Pt(14)

    summary_data = [
        ("1", "직접인건비", _fmt_won(labor_total), f"{cost_standard} {year}년 노임단가 적용"),
        ("2", "직접경비", _fmt_won(expense_total), "장비, SW, 출장 등"),
        ("3", "간접비", _fmt_won(overhead_total), f"직접인건비 × {overhead_rate*100 if overhead_rate < 2 else (overhead_rate-1)*100:.0f}%"),
        ("4", "기술료", _fmt_won(profit_total), f"(직접인건비+간접비) × {profit_rate*100:.0f}%"),
    ]

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 헤더
    hdr = table.rows[0]
    for i, txt in enumerate(["No", "비목", "금액 (원)", "산출 근거"]):
        _set_cell(hdr.cells[i], txt, bold=True, align="center", size=10)
    _shade_row(hdr, "D9E2F3")

    for no, name, amount, basis in summary_data:
        row = table.add_row()
        _set_cell(row.cells[0], no, align="center")
        _set_cell(row.cells[1], name)
        _set_cell(row.cells[2], amount, align="right")
        _set_cell(row.cells[3], basis)

    # 소계
    row = table.add_row()
    _set_cell(row.cells[0], "", align="center")
    _set_cell(row.cells[1], "소계", bold=True)
    _set_cell(row.cells[2], _fmt_won(subtotal), bold=True, align="right")
    _set_cell(row.cells[3], "")
    _shade_row(row, "E8E8E8")

    # 부가세
    row = table.add_row()
    _set_cell(row.cells[0], "5", align="center")
    _set_cell(row.cells[1], "부가가치세")
    _set_cell(row.cells[2], _fmt_won(vat), align="right")
    _set_cell(row.cells[3], "소계 × 10%")

    # 합계
    row = table.add_row()
    _set_cell(row.cells[0], "", align="center")
    _set_cell(row.cells[1], "합계", bold=True)
    _set_cell(row.cells[2], _fmt_won(grand_total), bold=True, align="right")
    _set_cell(row.cells[3], "")
    _shade_row(row, "D9E2F3")

    # 열 너비 설정
    widths = [Cm(1.5), Cm(4), Cm(4), Cm(7)]
    for row_obj in table.rows:
        for i, w in enumerate(widths):
            row_obj.cells[i].width = w

    doc.add_paragraph()

    # ═══════════════════════════════════════
    # 3. 직접인건비 상세
    # ═══════════════════════════════════════
    h = doc.add_heading("2. 직접인건비 산출 내역", level=1)
    h.runs[0].font.size = Pt(14)

    p = doc.add_paragraph()
    p.add_run(f"적용 기준: {cost_standard} {year}년 소프트웨어기술자 노임단가").font.size = Pt(9)

    if labor_breakdown:
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        hdr = table.rows[0]
        for i, txt in enumerate(["구분 (등급)", "월 단가 (원)", "투입 MM", "소계 (원)", "비고"]):
            _set_cell(hdr.cells[i], txt, bold=True, align="center", size=10)
        _shade_row(hdr, "D9E2F3")

        for item in labor_breakdown:
            grade = item.get("grade", item.get("role", ""))
            rate = item.get("monthly_rate", 0)
            mm = item.get("mm", item.get("person_months", 0))
            sub = item.get("subtotal", item.get("amount", int(rate * mm)))
            role = item.get("role", "")

            row = table.add_row()
            _set_cell(row.cells[0], grade)
            _set_cell(row.cells[1], _fmt_won(rate), align="right")
            _set_cell(row.cells[2], f"{mm:.1f}", align="center")
            _set_cell(row.cells[3], _fmt_won(sub), align="right")
            _set_cell(row.cells[4], role)

        # 합계
        row = table.add_row()
        _set_cell(row.cells[0], "합계", bold=True)
        _set_cell(row.cells[1], "")
        total_mm = sum(item.get("mm", item.get("person_months", 0)) for item in labor_breakdown)
        _set_cell(row.cells[2], f"{total_mm:.1f}", bold=True, align="center")
        _set_cell(row.cells[3], _fmt_won(labor_total), bold=True, align="right")
        _set_cell(row.cells[4], "")
        _shade_row(row, "E8E8E8")

        widths = [Cm(3.5), Cm(3.5), Cm(2.5), Cm(3.5), Cm(3.5)]
        for row_obj in table.rows:
            for i, w in enumerate(widths):
                row_obj.cells[i].width = w
    else:
        doc.add_paragraph("(인력 투입 상세 미산출 — STEP 3 실행 전)", style="List Bullet")

    doc.add_paragraph()

    # ═══════════════════════════════════════
    # 4. 직접경비 상세
    # ═══════════════════════════════════════
    h = doc.add_heading("3. 직접경비 산출 내역", level=1)
    h.runs[0].font.size = Pt(14)

    if expense_items:
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        hdr = table.rows[0]
        for i, txt in enumerate(["항목", "금액 (원)", "산출 근거", "비고"]):
            _set_cell(hdr.cells[i], txt, bold=True, align="center", size=10)
        _shade_row(hdr, "D9E2F3")

        for item in expense_items:
            row = table.add_row()
            _set_cell(row.cells[0], item.get("name", ""))
            _set_cell(row.cells[1], _fmt_won(item.get("amount", 0)), align="right")
            _set_cell(row.cells[2], item.get("basis", ""))
            _set_cell(row.cells[3], "")

        row = table.add_row()
        _set_cell(row.cells[0], "합계", bold=True)
        _set_cell(row.cells[1], _fmt_won(expense_total), bold=True, align="right")
        _set_cell(row.cells[2], "")
        _set_cell(row.cells[3], "")
        _shade_row(row, "E8E8E8")

        widths = [Cm(4), Cm(3.5), Cm(6), Cm(3)]
        for row_obj in table.rows:
            for i, w in enumerate(widths):
                row_obj.cells[i].width = w
    else:
        doc.add_paragraph("(직접경비 항목 없음)", style="List Bullet")

    doc.add_paragraph()

    # ═══════════════════════════════════════
    # 5. 간접비·기술료
    # ═══════════════════════════════════════
    h = doc.add_heading("4. 간접비 및 기술료", level=1)
    h.runs[0].font.size = Pt(14)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr = table.rows[0]
    for i, txt in enumerate(["비목", "산출 방법", "금액 (원)"]):
        _set_cell(hdr.cells[i], txt, bold=True, align="center", size=10)
    _shade_row(hdr, "D9E2F3")

    indirect_items = [
        ("간접비", f"직접인건비({_fmt_won(labor_total)}) × {overhead_rate*100 if overhead_rate < 2 else (overhead_rate-1)*100:.0f}%", _fmt_won(overhead_total)),
        ("기술료", f"(직접인건비+간접비) × {profit_rate*100:.0f}%", _fmt_won(profit_total)),
    ]
    for name, method, amount in indirect_items:
        row = table.add_row()
        _set_cell(row.cells[0], name)
        _set_cell(row.cells[1], method)
        _set_cell(row.cells[2], amount, align="right")

    widths = [Cm(3), Cm(10), Cm(3.5)]
    for row_obj in table.rows:
        for i, w in enumerate(widths):
            row_obj.cells[i].width = w

    doc.add_paragraph()

    # ═══════════════════════════════════════
    # 6. 산출 근거 (Budget Narrative)
    # ═══════════════════════════════════════
    if budget_narrative:
        h = doc.add_heading("5. 산출 근거", level=1)
        h.runs[0].font.size = Pt(14)

        for item in budget_narrative:
            if isinstance(item, dict):
                cost_item = item.get("cost_item", "")
                activity = item.get("linked_activity", "")
                justification = item.get("justification", "")
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"{cost_item}")
                run.bold = True
                run.font.size = Pt(10)
                if activity:
                    p.add_run(f" — {activity}").font.size = Pt(9)
                if justification:
                    detail = doc.add_paragraph(f"  {justification}")
                    detail.runs[0].font.size = Pt(9)
                    detail.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── 파일 저장 ──
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
