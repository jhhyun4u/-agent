from pathlib import Path

from docx import Document
from docx.shared import Pt


def build_docx(sections: dict, output_path: Path, project_name: str = "용역 제안서") -> Path:
    doc = Document()

    title = doc.add_heading(f"{project_name} 제안서", level=0)
    title.runs[0].font.size = Pt(24)

    for i, (heading, body) in enumerate(sections.items(), start=1):
        doc.add_heading(f"{i}. {heading}", level=1)
        if isinstance(body, str):
            for paragraph in body.split("\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
