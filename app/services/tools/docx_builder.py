"""
DOCX 빌더 (§9-1)

케이스 A(자유 양식) / 케이스 B(서식 지정) 분기.
중간 버전 + 최종 버전 모두 생성 가능.
"""

import asyncio
import io
import logging
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


# ── 레거시 호환 (기존 코드에서 사용) ──

def build_docx_legacy(sections: dict, output_path: Path, project_name: str = "용역 제안서") -> Path:
    """레거시 DOCX 빌더 (기존 인터페이스 유지)."""
    doc = Document()
    title = doc.add_heading(f"{project_name} 제안서", level=0)
    title.runs[0].font.size = Pt(24)

    for i, (heading, body) in enumerate(sections.items(), start=1):
        doc.add_heading(f"{i}. {heading}", level=1)
        if isinstance(body, str):
            for paragraph in body.split("\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ── v3.4 DOCX 빌더 (§9-1) ──

async def build_docx(sections: list, rfp=None, proposal_name: str = "") -> bytes:
    """제안서 DOCX 생성. 케이스 A/B에 따라 분기. bytes 반환."""
    rfp_dict = {}
    if rfp:
        rfp_dict = rfp.model_dump() if hasattr(rfp, "model_dump") else (rfp if isinstance(rfp, dict) else {})

    case_type = rfp_dict.get("case_type", "A")

    if case_type == "B" and rfp_dict.get("format_template", {}).get("structure"):
        return await asyncio.to_thread(_build_from_template, sections, rfp_dict["format_template"]["structure"], proposal_name)
    else:
        return await asyncio.to_thread(_build_freeform, sections, rfp_dict, proposal_name)


def _build_freeform(sections: list, rfp_dict: dict, proposal_name: str) -> bytes:
    """케이스 A: 자유 양식 DOCX."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(11)

    # 표지
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(proposal_name or rfp_dict.get("project_name", "제안서"))
    run.bold = True
    run.font.size = Pt(24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("기술제안서")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    if rfp_dict.get("client"):
        client_p = doc.add_paragraph()
        client_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        client_p.add_run(f"\n\n{rfp_dict['client']}").font.size = Pt(14)

    doc.add_page_break()

    # 목차
    doc.add_heading("목 차", level=1)
    for i, s in enumerate(sections, 1):
        s_dict = s.model_dump() if hasattr(s, "model_dump") else (s if isinstance(s, dict) else {})
        title_text = s_dict.get("title", s_dict.get("section_id", f"섹션 {i}"))
        doc.add_paragraph(f"{i}. {title_text}", style="List Number")
    doc.add_page_break()

    # 본문
    for i, s in enumerate(sections, 1):
        s_dict = s.model_dump() if hasattr(s, "model_dump") else (s if isinstance(s, dict) else {})
        title_text = s_dict.get("title", s_dict.get("section_id", f"섹션 {i}"))
        content = s_dict.get("content", "")

        doc.add_heading(f"{i}. {title_text}", level=1)
        _render_markdown_content(doc, content)
        doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _build_from_template(sections: list, structure: dict, proposal_name: str) -> bytes:
    """케이스 B: 서식 구조 보존 DOCX."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(proposal_name or "기술제안서")
    run.bold = True
    run.font.size = Pt(24)
    doc.add_page_break()

    # 서식 구조대로 배치
    section_map = {}
    for s in sections:
        s_dict = s.model_dump() if hasattr(s, "model_dump") else (s if isinstance(s, dict) else {})
        section_map[s_dict.get("section_id", "")] = s_dict

    for section_id, template in structure.items():
        s_dict = section_map.get(section_id, {})
        title_text = template.get("title", section_id) if isinstance(template, dict) else section_id
        content = s_dict.get("content", "(작성 필요)")

        doc.add_heading(title_text, level=1)
        _render_markdown_content(doc, content)
        doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _render_markdown_content(doc: Document, content: str):
    """Markdown → DOCX 변환 (기본 수준)."""
    if not content:
        return

    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            doc.add_paragraph(text, style="List Number")
        elif stripped.startswith("|") and stripped.endswith("|"):
            doc.add_paragraph(stripped, style="Normal")
        else:
            p = doc.add_paragraph()
            parts = re.split(r"\*\*(.*?)\*\*", stripped)
            for j, part in enumerate(parts):
                run = p.add_run(part)
                if j % 2 == 1:
                    run.bold = True
